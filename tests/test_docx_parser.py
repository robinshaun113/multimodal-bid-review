from docx import Document

from docx_parser import parse_docx


def test_parser_preserves_paragraph_table_order_and_section(tmp_path):
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_heading("技术要求", level=1)
    doc.add_paragraph("第一段要求")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = "SLA"
    table.cell(1, 1).text = "不低于99.9%"
    doc.add_paragraph("表后说明")
    doc.save(path)

    parsed = parse_docx(path, tmp_path / "images")
    chunks = parsed["text_chunks"]
    assert [c["kind"] for c in chunks] == [
        "paragraph",
        "paragraph",
        "table",
        "paragraph",
    ]
    assert chunks[2]["section"] == "技术要求"
    assert "SLA | 不低于99.9%" in chunks[2]["text"]
    assert chunks[1]["block_index"] < chunks[2]["block_index"] < chunks[3]["block_index"]
